import sys
import os
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches

from models.tables import Hymn, HymnContent, ContentLine
from core.exceptions import HimnarioGeneratorException, DatabaseError, HymnNotFoundError

# Add the parent directory of the modules to the Python path
# sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

# Assuming HymnDocumentGenerator is now part of the backend or can be adapted
# from modules.generator_hymnary.hymn_document_generator import HymnDocumentGenerator

async def generate_hymnary_docx(db: Session, hymn_ids: list[int], file_name: str) -> str:
    try:
        # Fetch hymns from the database
        hymns = db.query(Hymn).filter(Hymn.id.in_(hymn_ids)).all()
        if not hymns:
            # Usar -1 como id inválido para indicar que no se encontró ningún himno
            raise HymnNotFoundError(hymn_id=-1)

        document = Document()
        document.add_heading('Himnario Generado', level=1)

        for hymn in hymns:
            document.add_heading(f'{hymn.hymn_number}. {hymn.title}', level=2)
            
            # Assuming content and lines are loaded via joinedload in hymn_service.get_hymn
            # If not, you might need to load them here or adjust the query.
            for content_item in sorted(hymn.content, key=lambda x: x.content_order):
                if content_item.content_type == 'estrofa':
                    document.add_paragraph(f'Estrofa {content_item.stanza_number}')
                elif content_item.content_type == 'coro':
                    document.add_paragraph('Coro')
                
                for line in sorted(content_item.lines, key=lambda x: x.line_order):
                    document.add_paragraph(line.line_text)

            document.add_page_break()

        # Ensure output directory exists
        output_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'modules', 'generator_hymnary', 'output')
        os.makedirs(output_dir, exist_ok=True)

        output_filepath = os.path.join(output_dir, file_name)
        document.save(output_filepath)

        return output_filepath

    except HymnNotFoundError:
        raise # Re-raise specific error
    except Exception as e:
        raise HimnarioGeneratorException(detail=f"Failed to generate DOCX document: {e}")
